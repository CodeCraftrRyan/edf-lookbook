from flask import Flask, render_template, request, send_file, redirect, url_for, session, flash
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash 
from user_models import db, User  # type: ignore
import os 
import pandas as pd
from docx import Document
from docx.shared import Inches, RGBColor, Pt 
from docx.enum.table import WD_ALIGN_VERTICAL
from PIL import Image as PILImage
from datetime import datetime
from flask_mail import Mail, Message 
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
import csv

# Login and registration features are temporarily disabled for testing

app = Flask(__name__, instance_relative_config=True)
# Email configuration
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = 'your_email@gmail.com'   
app.config['MAIL_PASSWORD'] = 'your_email_password'       
app.config['MAIL_DEFAULT_SENDER'] = 'your_email@gmail.com' 

mail = Mail(app)

db_path = os.path.join(app.instance_path, 'users.db')
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{db_path}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db.init_app(app)
app.secret_key = 'dolefoundation'  

#login_manager = LoginManager()
#login_manager.init_app(app)
#login_manager.login_view = 'login' 

# Folders 
app.config['UPLOAD_FOLDER'] = "uploads"
IMAGE_FOLDER = "static/images"
TEMP_IMAGE_FOLDER = "processed_images"

#Create folders if needed
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(TEMP_IMAGE_FOLDER, exist_ok=True)

#Flask-Login user loader
#@login_manager.user_loader
#def load_user(user_id):
#    return User.query.get(int(user_id))

#Login route 
# @app.route('/login', methods=['GET', 'POST'])
# def login():
#   if request.method == 'POST':
#         username = request.form.get('username')
#         password = request.form.get('password')
#         user = User.query.filter_by(username=username).first()
#         if user and user.check_password(password):
#             login_user(user)
#             return redirect(url_for('upload_csv'))
#         else:
#             flash("Invalid username or password.")
#     return render_template('login.html')

'''
@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form.get("username")
        email = request.form.get("email")
        password = request.form.get("password")
        
        # ✅ Restrict to EDF domain
        if not email.endswith("@elizabethdolefoundation.org"):
            flash("Only EDF staff emails are allowed.")
            return redirect(url_for("register"))

        # Check if username already exists in the database
        existing_user = User.query.filter_by(username=username).first()
        if existing_user:
            flash("Username already taken. Please choose another.")
            return redirect(url_for("register"))
        
        if User.query.filter_by(email=email).first():
            flash("Email already in use.")
            return redirect(url_for("register"))

        # Create and store new user
        new_user = User(username=username, email=email)
        new_user.set_password(password)  # securely hash password
        db.session.add(new_user)
        db.session.commit()

        flash("Account created! Please log in.")
        return redirect(url_for("login"))

    return render_template("register.html")
'''

'''
@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        username = request.form['username']
        user = User.query.filter_by(username=username).first()

        if user:
            reset_link = f"http://localhost:5001/reset_password/{user.id}"
            msg = Message("Reset Your Password",
                          recipients=[user.email])
            msg.body = f"Hi {user.username},\n\nClick here to reset your password: {reset_link}\n\nIf you didn’t ask for this, you can ignore it."
            mail.send(msg)
            flash("✅ Reset email sent!")
        else:
            flash("⚠️ No user found with that username.")

        return redirect(url_for('login'))

    return render_template("forgot_password.html")
'''
# Utility function to set column width
from docx.shared import Inches

def set_column_width(table, column_idx, width_in_inches):
    for cell in table.columns[column_idx].cells:
        cell.width = Inches(width_in_inches)

def resolve_image_by_id(image_key):
    """
    Given an image base key (usually lookbookID), try common extensions in IMAGE_FOLDER.
    Returns a full path if found, else empty string.
    """
    if not image_key:
        return ""
    base, ext = os.path.splitext(image_key)
    # If an extension is already present, check directly
    if ext:
        candidate = os.path.join(IMAGE_FOLDER, image_key)
        return candidate if os.path.exists(candidate) else ""
    # Otherwise, try common extensions in priority order
    for e in [".jpg", ".jpeg", ".png", ".JPG", ".JPEG", ".PNG"]:
        candidate = os.path.join(IMAGE_FOLDER, image_key + e)
        if os.path.exists(candidate):
            return candidate
    return ""

#Main upload + PDF generation route
@app.route("/", methods=["GET", "POST"])
#@login_required
def upload_csv():
    if request.method == "POST":
        file = request.files["file"]
        if file and file.filename.endswith(".csv"):
            csv_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(csv_path)

            # Load CSV with robust handling, including Salesforce "CSV Unicode (UTF-16)" exports
            read_kwargs = dict(sep=None, engine="python")  # allow auto-detect of comma vs tab
            try:
                df = pd.read_csv(csv_path, encoding="utf-8-sig", **read_kwargs)
            except UnicodeDecodeError:
                # Try UTF-16 variants (Salesforce "CSV Unicode" often uses BE; pandas will auto-detect with BOM)
                for enc in ("utf-16", "utf-16-be", "utf-16-le", "cp1252", "latin1"):
                    try:
                        df = pd.read_csv(csv_path, encoding=enc, **read_kwargs)
                        break
                    except UnicodeDecodeError:
                        df = None
                if df is None:
                    # Final fallback without explicit encoding (let pandas guess)
                    df = pd.read_csv(csv_path, **read_kwargs)
            df.fillna("", inplace=True)

            # Normalize common Windows control chars to smart punctuation (if any slipped through)
            def _normalize_text(val):
                if not isinstance(val, str):
                    return val
                return (
                    val.replace("\u0091", "\u2018")   # ‘
                       .replace("\u0092", "\u2019")   # ’
                       .replace("\u0093", "\u201C")   # “
                       .replace("\u0094", "\u201D")   # ”
                       .replace("\u0096", "\u2013")   # –
                       .replace("\u0097", "\u2014")   # —
                )
            df = df.applymap(_normalize_text)

            output_path = "lookbook output.docx"
            doc = Document()
            # Set global font to Times New Roman, 11pt
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(11)
            
            # Add a logo on the title page if available
            logo_path = os.path.join("static", "images", "logo.png")
            if os.path.exists(logo_path):
                try:
                    doc.add_picture(logo_path, width=Inches(1.7))
                except Exception as e:
                    doc.add_paragraph("[Logo error]")
            
            doc.add_heading("Elizabeth Dole Foundation Look-Book", 0)
            doc.add_paragraph("Generated on: " + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            doc.add_page_break()
            

            from docx.shared import RGBColor

            # Track image selection by lookbookID so repeated IDs reuse the same image
            id_to_image = {}

            for index, row in df.iterrows():
                # New canonical headers with graceful fallback to legacy names
                name = (row.get("Full Name") or row.get("Name") or "").strip()
                title_text = (row.get("Primary Affiliation Role (Lookbook)") or row.get("Title") or "").strip()
                company = (row.get("Primary Organization (Lookbook)") or row.get("Company") or "").strip()
                additional = (row.get("Additional") or "").strip()
                # Track ID for cross-row consistency
                lookbook_id = (row.get("lookbookID") or row.get("LookbookID") or "").strip()
                # Long text: new header "Bio/About Me" or legacy "Description"/"Bio"
                description = ((row.get("Bio/About Me") or row.get("Description") or row.get("Bio") or "")).strip()
                # Images use the Lookbook ID as the filename base (e.g., LB-0001.jpg/.png)
                image_key = lookbook_id
                if lookbook_id:
                    # Cache the image key for this ID so repeats stay consistent
                    if lookbook_id not in id_to_image:
                        id_to_image[lookbook_id] = image_key
                    image_key = id_to_image[lookbook_id]
                image_path = resolve_image_by_id(image_key)
                
                # Build header table (image + name/title/company)
                table = doc.add_table(rows=1, cols=3)
                table.autofit = False
                # Set column widths
                set_column_width(table, 0, 1.6)   # Image column
                set_column_width(table, 1, 4.2)   # Text column (wider for long titles)
                set_column_width(table, 2, 0.6)   # Spacer column

                # Column 1: Image
                img_cell = table.cell(0, 0)
                img_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                if os.path.exists(image_path):
                    try:
                        img_para = img_cell.paragraphs[0]
                        img_para.paragraph_format.space_after = 0
                        run = img_para.add_run()
                        run.add_picture(image_path, width=Inches(1.7))
                    except Exception as e:
                        img_cell.text = "[Image error]"
                else:
                    img_cell.text = f"[Image not found for ID: {lookbook_id}]"

                # Column 2: Text content (name + role)
                text_cell = table.cell(0, 1)
                name_para = text_cell.paragraphs[0]
                name_para.style = 'Heading 2'
                name_para.add_run(name)

                info_para = text_cell.add_paragraph()
                info_run = info_para.add_run(f"{company}, {title_text}")
                info_para.add_run("\n")
                if additional:
                    add_run = info_para.add_run(additional)
                    add_run.italic = True

                # Add a full-width Bio section below the table to allow long text across pages
                bio_title = doc.add_paragraph()
                bio_title_run = bio_title.add_run("Bio")
                bio_title_run.bold = True
                bio_title.paragraph_format.space_before = Pt(6)
                bio_title.paragraph_format.space_after = Pt(2)

                bio_para = doc.add_paragraph()
                bio_para.paragraph_format.space_after = Pt(12)
                bio_para.paragraph_format.line_spacing = 1.15
                bio_run = bio_para.add_run(description)

                doc.add_paragraph("")  # spacer between entries

            doc.save(output_path)
            return send_file(output_path, as_attachment=True)

    #GET request
    return render_template("upload.html")

#@app.route("/logout", methods=["GET"]) removed for testing
#@login_required temporarily removed for testing
#    logout_user()
#    flash("You’ve been logged out.")
#    return redirect(url_for("login"))

@app.route('/download-template')
def download_template():
    from io import StringIO
    import csv

    output = BytesIO()
    writer = csv.writer(output)
    
    import io
    text_stream = io.TextIOWrapper(output, encoding='utf-8')
    writer = csv.writer(text_stream)

    # New canonical headers used by the Lookbook
    writer.writerow(["lookbookID", "Full Name", "Primary Affiliation Role (Lookbook)", "Primary Organization (Lookbook)", "Bio/About Me"])

    # Example row using the new schema
    writer.writerow([
        "LB-0001",
        "Jane Doe",
        "Senior Advisor",
        "Elizabeth Dole Foundation",
        "Jane has dedicated over a decade to supporting caregivers across the country."
    ])

    text_stream.flush()
    output.seek(0)

    return send_file(
        output,
        mimetype='text/csv',
        as_attachment=True,
        download_name='LOOKBOOK_TEMPLATE.csv'
    )

#Run the app
if __name__ == '__main__':
    app.run(debug=True)